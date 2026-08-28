# In[0.1]: Install packages
# This script was originally written to run in notebook cells
# (Jupyter/IPython). To run it as a .py script, install the dependencies
# from the terminal first:
#   pip install pandas numpy seaborn matplotlib plotly scipy statsmodels
#   pip install scikit-learn statstests python-docx

# In[0.2]: Import packages
import warnings

import pandas as pd  # dataframe-based data manipulation
import numpy as np  # mathematical operations
import seaborn as sns  # data visualization
import matplotlib.pyplot as plt  # data visualization
import statsmodels.api as sm  # model estimation
import statsmodels.formula.api as smf  # model estimation (binary logistic and count models)
import plotly.graph_objects as go  # 3D charts
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from statstests.process import stepwise  # stepwise procedure

warnings.filterwarnings('ignore')

# %% LOADING
dados = pd.read_excel('703pacientes.xlsx')
dados.info()

# %% Frequency table
colunas = [
    'Vacinado',
    'Sexo',
    'Estado_Civil_1',
    'Estado_Civil_2',
    'Idade_cat_1',
    'Idade_cat_2',
    'Idade_cat_3',
    'Óbito',
    'Trombose',
    'Sepse', 'SRAG',
    'Choque'
    'Grau_Instrucao_1',
    'Grau_Instrucao_2',
    'UF',
    'Município',
    'Fabricante',
    'Nenhuma_dose',
    'Aztrazeneca',
    'Pfizer',
    'Coronavac_Butantan',
    'Janssen',
    'Prob_Card',
    'CP',
    'Diabetes',
    'SRAG',
    'Choques',
    'Prob_neurol',
    'Prob_Hemat',
    'Cancer',
    'Prob_Resp',
    'Prob_Infec',
    'Prob_Metab',
    'Prob_TGI',
    'Prob_Hep',
    'Prob_Hid_Elet',
    'Prob_AI_Infla',
    'Febre',
    'Traumatismo',
    'Covid_critica',
    'Prob_Renal',
    'LRA',
]
n_total = 720

# Create the Word document
doc = Document()
for coluna in colunas:
    freq_absoluta = dados[coluna].value_counts().rename('Absolute Frequency')
    freq_relativa = (dados[coluna].value_counts() / n_total * 100).round(2).rename('Relative Frequency (%)')
    tabela = pd.concat([freq_absoluta, freq_relativa], axis=1).reset_index()
    tabela.columns = [coluna, 'Absolute Frequency', 'Relative Frequency (%)']

    # Table title (ABNT format)
    titulo = doc.add_paragraph()
    run = titulo.add_run(f"Table - Frequency of variable {coluna}")
    run.bold = True
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo.space_after = Pt(12)

    # Create the table in Word
    table = doc.add_table(rows=1, cols=len(tabela.columns))
    table.style = 'Table Grid'

    # Table header
    for i, heading in enumerate(tabela.columns):
        cell = table.cell(0, i)
        cell.text = heading
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add table rows
    for idx, row in tabela.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Spacing after the table
    doc.add_paragraph()

# Save the Word file
nome_arquivo = 'frequencias_vaccine.docx'
doc.save(nome_arquivo)

# %% Frequency table
colunas = [
    'Óbito',
    'Trombose',
    'Sepse',
    'SRAG',
    'Choque',
]
n_total = 720

# Create the Word document
doc = Document()
for coluna in colunas:
    freq_absoluta = dados[coluna].value_counts().rename('Absolute Frequency')
    freq_relativa = (dados[coluna].value_counts() / n_total * 100).round(2).rename('Relative Frequency (%)')
    tabela = pd.concat([freq_absoluta, freq_relativa], axis=1).reset_index()
    tabela.columns = [coluna, 'Absolute Frequency', 'Relative Frequency (%)']

    # Table title (ABNT format)
    titulo = doc.add_paragraph()
    run = titulo.add_run(f"Table - Frequency of variable {coluna}")
    run.bold = True
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo.space_after = Pt(12)

    # Create the table in Word
    table = doc.add_table(rows=1, cols=len(tabela.columns))
    table.style = 'Table Grid'

    # Table header
    for i, heading in enumerate(tabela.columns):
        cell = table.cell(0, i)
        cell.text = heading
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add table rows
    for idx, row in tabela.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Spacing after the table
    doc.add_paragraph()

# Save the Word file
nome_arquivo = 'frequencias_pacientes_covid_critic.docx'
doc.save(nome_arquivo)

# %% Multiple simple logistic regressions
# ================================
# 1. LIST OF INDEPENDENT VARIABLES
# ================================
variaveis = [
    'Óbito',
    'SRAG',
    'Prob_Card',
    'CP',
    'Diabetes',
    'SRAG',
    'Choques',
    'Prob_neurol',
    'Prob_Hemat',
    'Cancer',
    'Prob_Resp',
    'Prob_Infec',
    'Prob_Metab',
    'Prob_TGI',
    'Prob_Hep',
    'Prob_Hid_Elet',
    'Prob_AI_Infla',
    'Febre',
    'Traumatismo',
    'Prob_Renal',
    'LRA',
    'Vacinado',
    'Estado_Civil_1',
    'Estado_Civil_2',
    'Idade_cat_1',
    'Idade_cat_2',
    'Idade_cat_3',
    'Fabricante',
    'Nenhuma_dose',
    'Aztrazeneca',
    'Grau_Instrucao_1',
    'Grau_Instrucao_2',
]

# Dependent variable
y = 'Óbito'

# ================================
# 2. LIST TO STORE THE RESULTS
# ================================
resultados = []

# ================================
# 3. RUN ONE REGRESSION PER VARIABLE
# ================================
for var in variaveis:
    formula = f"{y} ~ {var}"
    modelo = smf.glm(formula=formula, data=dados, family=sm.families.Binomial()).fit()
    beta = modelo.params[var]
    p = modelo.pvalues[var]
    ic_low, ic_high = modelo.conf_int().loc[var]
    OR = np.exp(beta)
    OR_low = np.exp(ic_low)
    OR_high = np.exp(ic_high)
    resultados.append({
        'Variable': var,
        'Beta': beta,
        'IC95_low': ic_low,
        'IC95_high': ic_high,
        'p-value': p,
        'OR': OR,
        'OR_low': OR_low,
        'OR_high': OR_high,
    })

# ================================
# 4. BUILD THE FINAL DATAFRAME
# ================================
pd.set_option('display.max_rows', None)  # show all rows
pd.set_option('display.max_columns', None)  # show all columns
pd.set_option('display.width', None)  # avoid line wrapping
pd.set_option('display.max_colwidth', None)  # avoid truncating content
pd.set_option('display.float_format', '{:.6f}'.format)
df_resultados = pd.DataFrame(resultados)
print(df_resultados)

# %% Word table with the logistic regression results
document = Document()
document.add_heading("Logistic Regression Results", level=1)

# Create table with the correct number of columns
tabela = document.add_table(rows=1, cols=len(df_resultados.columns))

# Fill header
hdr_cells = tabela.rows[0].cells
for i, col in enumerate(df_resultados.columns):
    hdr_cells[i].text = str(col)

# Fill dataframe rows
for idx, row in df_resultados.iterrows():
    row_cells = tabela.add_row().cells
    for i, col in enumerate(df_resultados.columns):
        valor = row[col]
        # round numeric values
        if isinstance(valor, float):
            valor = round(valor, 4)
        row_cells[i].text = str(valor)

# Save file
document.save("resultados_regressoes_binarias2.docx")
print("Word file generated successfully!")

# %% Logistic regression with all independent variables
modelo_rl = smf.glm(
    formula='Óbito ~ Vacinado + Pfizer + Sexo + Nenhuma_dose + Prob_Card + '
            'Aztrazeneca + CP + UF + Município + Diabetes + SRAG + Pfizer + Choques + '
            'Prob_neurol + Janssen + Prob_Hemat + Coronavac_Butantan + Cancer + '
            'Prob_Resp + Prob_Infec + Prob_Metab + Prob_TGI + Prob_Hep + Prob_Hid_Elet + '
            'Prob_AI_Infla + Febre + Outros + Traumatismo + COVID_CRÍTICA + Prob_Renal + '
            'LRA + Dias_permanência + Estado_Civil_1 + Estado_Civil_2 + Idade_cat_1 + '
            'Idade_cat_2 + Idade_cat_3 + Grau_Instrucao_1 + Grau_Instrucao_2',
    data=dados, family=sm.families.Binomial(),
).fit()
modelo_rl.summary()

########################## ADDITIONAL PART #################################
# OR
or_values = np.exp(modelo_rl.params)

# 95% CI of the coefficients -> also transform into OR
conf = modelo_rl.conf_int()
conf['OR_inf'] = np.exp(conf[0])
conf['OR_sup'] = np.exp(conf[1])

# Final organized table
resultado = pd.DataFrame({
    'Coef (β)': modelo_rl.params,
    'OR': or_values,
    'IC_inf': conf['OR_inf'],
    'IC_sup': conf['OR_sup'],
    'p-value': modelo_rl.pvalues,
})
pd.set_option('display.float_format', '{:.6f}'.format)
print(resultado)

# %%
# Model estimation via the Stepwise procedure
step_modelo = stepwise(modelo_rl, pvalue_limit=0.05)

# %% swapping "vacinado" for "vacinas" (both can't be in the same
# model at once due to multicollinearity).
dados = pd.get_dummies(dados, columns=['Vacinas'], dtype=int, drop_first=True)
dados.info()

# %% Logistic regression with all independent variables
modelo_rl_vacinado = smf.glm(
    formula='Óbito ~ Vacinas_1 + Vacinas_2 + Vacinas_3 + Vacinas_4 + Sexo + '
            'Prob_Card + CP + Diabetes + SRAG + Choques + Prob_neurol + Prob_Hemat + '
            'Cancer + Prob_Resp + Prob_Infec + Prob_Metab + Prob_TGI + Prob_Hep + '
            'Prob_Hid_Elet + Prob_AI_Infla + Febre + Outros + Traumatismo + '
            'COVID_CRÍTICA + Prob_Renal + LRA + Dias_permanência + Estado_Civil_1 + '
            'Estado_Civil_2 + Idade_cat_1 + Idade_cat_2 + Idade_cat_3 + '
            'Grau_Instrucao_1 + Grau_Instrucao_2',
    data=dados, family=sm.families.Binomial(),
).fit()
modelo_rl_vacinado.summary()

# Coefficients
betas = modelo_rl_vacinado.params

# Odds Ratio
or_values = np.exp(betas)

# Confidence interval (β)
conf = modelo_rl_vacinado.conf_int()

# CI transformed into OR
or_low = np.exp(conf[0])
or_high = np.exp(conf[1])

# Building the final table
resultado_vacinas = pd.DataFrame({
    'Variable': betas.index,
    'Beta': betas.values,
    'OR': or_values.values,
    'OR_low': or_low.values,
    'OR_high': or_high.values,
    'p-value': modelo_rl_vacinado.pvalues.values,
})

# Optional formatting
pd.set_option('display.float_format', '{:.6f}'.format)
print(resultado_vacinas.to_string(index=False))

# %%
# Model estimation via the Stepwise procedure
step_modelo = stepwise(modelo_rl_vacinado, pvalue_limit=0.05)
